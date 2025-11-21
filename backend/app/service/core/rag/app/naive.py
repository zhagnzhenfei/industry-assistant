#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
from tika import parser
from io import BytesIO
from docx import Document
from timeit import default_timer as timer
import re
from service.core.deepdoc.parser.pdf_parser import PlainParser
from service.core.rag.nlp import rag_tokenizer, naive_merge, tokenize_table, tokenize_chunks, find_codec, concat_img, \
    naive_merge_docx, tokenize_chunks_docx
from service.core.deepdoc.parser import PdfParser, ExcelParser, DocxParser, HtmlParser, JsonParser, MarkdownParser, TxtParser
from service.core.rag.app.table import chunk as table_chunk
from service.core.rag.utils import num_tokens_from_string
from PIL import Image
from functools import reduce
from markdown import markdown
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError


class Docx(DocxParser):
    """
    Word文档解析器
    处理.docx文件，提取文本内容、图片和表格
    """
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        """
        从Word段落中提取图片
        返回PIL图像对象或None
        """
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            logging.info("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception:
            return None

    def __clean(self, line):
        """
        清理文本行，移除特殊字符
        """
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        """
        解析Word文档，提取文本、图片和表格
        返回处理后的文本行和表格数据
        """
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    """
    PDF文档解析器
    处理.pdf文件，使用OCR识别文本、分析布局和提取表格
    """
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        """
        解析PDF文档，提取文本内容和表格
        使用OCR和布局分析技术
        """
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))
        tbls = self._extract_table_figure(True, zoomin, True, True)
        self._concat_downward()

        logging.info("layouts cost: {}s".format(timer() - first_start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


class Markdown(MarkdownParser):
    """
    Markdown文档解析器
    处理.md文件，提取文本内容和表格
    """
    def __call__(self, filename, binary=None):
        """
        解析Markdown文档，提取文本内容和表格
        将内容分割成适当大小的块
        """
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')
        sections = []
        tbls = []
        for sec in remainder.split("\n"):
            if num_tokens_from_string(sec) > 3 * self.chunk_token_num:
                sections.append((sec[:int(len(sec) / 2)], ""))
                sections.append((sec[int(len(sec) / 2):], ""))
            else:
                if sec.strip().find("#") == 0:
                    sections.append((sec, ""))
                elif sections and sections[-1][0].strip().find("#") == 0:
                    sec_, _ = sections.pop(-1)
                    sections.append((sec_ + "\n" + sec, ""))
                else:
                    sections.append((sec, ""))

        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        return sections, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
    文档分块处理函数
    根据文件类型选择合适的解析器处理文档内容
    将文档内容分割成适合RAG模型的块
    
    支持的文件格式：docx, pdf, excel, txt, md, html, json等
    将连续文本按照'delimiter'分割成小块
    然后将这些块合并成token数不超过'Max token number'的块
    """

    is_english = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tables = Docx()(filename, binary)
        res = tokenize_table(tables, doc, is_english)  # just for table

        callback(0.8, "Finish parsing.")
        st = timer()

        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        if kwargs.get("section_only", False):
            return chunks

        res.extend(tokenize_chunks_docx(chunks, doc, is_english, images))
        logging.info("naive_merge({}): {}".format(filename, timer() - st))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        if parser_config.get("layout_recognize", "DeepDOC") == "Plain Text":
            pdf_parser = PlainParser()
        sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page,
                                      callback=callback)
        res = tokenize_table(tables, doc, is_english)

    elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = ExcelParser()
        
        # 如果binary为None，从文件路径读取
        if binary is None:
            with open(filename, 'rb') as f:
                binary = f.read()
        
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.csv$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse CSV.")
        # 使用table.py中的CSV解析功能
        try:
            csv_chunks = table_chunk(filename, binary, from_page=from_page, to_page=to_page,
                                   lang=lang, callback=callback, **kwargs)
            res.extend(csv_chunks)
            callback(0.8, "CSV parsing completed.")
            return res
        except Exception as e:
            logging.error(f"CSV parsing failed: {e}")
            raise NotImplementedError(f"CSV parsing failed: {e}")

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tables = Markdown(int(parser_config.get("chunk_token_num", 128)))(filename, binary)
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        if doc_parsed.get('content', None) is not None:
            sections = doc_parsed['content'].split('\n')
            sections = [(_, "") for _ in sections if _]
            callback(0.8, "Finish parsing.")
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt, csv supported)")

    st = timer()
    chunks = naive_merge(
        sections, int(parser_config.get(
            "chunk_token_num", 128)), parser_config.get(
            "delimiter", "\n!?。；！？"))
    if kwargs.get("section_only", False):
        return chunks

    res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))
    logging.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
